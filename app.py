import streamlit as st
import pandas as pd
import numpy as np
import io
import json
import xlsxwriter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import docx.opc.constants
import time

# ==========================================
# PART 1: 配置区域 (已增强 'add_to_cart' 映射)
# ==========================================

COMMON_METRICS = {
    "spend": ["花费金额(USD)", "花费金额 （USD）", "花费金额 (USD)", "花费金额", "Amount Spent", "Cost"],
    "roas": ["广告花费回报 (ROAS) - 购物", "广告花费回报（ROAS）-购物", "ROAS", "Purchase ROAS", "Return on Ad Spend"],
    "purchases": ["购买次数", "成效数量", "成效", "Purchases", "Results", "Website Purchases"],
    "cpa": ["单次购买费用", "单次购物成本", "单次成效成本", "单次成效费用", "Cost per Purchase", "Cost per Result"],
    "ctr": ["链接点击率", "链接点击率（%)", "链接点击率（%）", "CTR", "Link CTR"],
    "cpm": ["千次展示费用", "CPM", "Cost per 1,000 Impressions"],
    "clicks": ["点击", "链接点击", "Clicks", "Link Clicks"],
    "impressions": ["曝光", "展示次数", "Impressions"],
    "purchase_value": ["购买价值", "购物价值", "Purchase Value", "Conversion Value"],
    "aov": ["单次购买价值", "单次购物价值"]
}

SHEET_MAPPINGS = {
    "整体数据": {
        **COMMON_METRICS,
        "date_range": ["时间范围", "Date Range", "Time"],
        "clicks_all": ["点击", "点击(全部)", "Clicks (All)"],
        "landing_page_views": ["落地页浏览量", "落地页", "Landing Page Views", "Landing"],
        # ✅ 修改点：增加了更多常见的加购列名别名
        "add_to_cart": ["加入购物车", "加购", "Add to Cart", "Website Adds to Cart", "网站加购", "Adds to Cart"], 
        "initiate_checkout": ["结账发起次数", "结账", "Initiate Checkout", "Website Initiated Checkouts", "网站结账发起"],
        "rate_click_to_lp": ["点击-落地页浏览转化率"],
        "rate_lp_to_atc": ["落地页浏览-加购转化率"],
        "rate_atc_to_ic": ["加购-结账转化率"],
        "rate_ic_to_pur": ["结账-购买转化率"]
    },
    "分时段数据": {
        **COMMON_METRICS,
        "date_range": ["时间范围", "Day", "Date"],
        "landing_page_views": ["落地页浏览量", "Landing Page Views"],
        # ✅ 修改点：确保这里包含【加入购物车】以及其他变体
        "add_to_cart": ["加入购物车", "加购", "Add to Cart", "Website Adds to Cart", "网站加购", "Adds to Cart"],
        "initiate_checkout": ["结账发起次数", "Initiate Checkout"],
        "rate_click_to_lp": ["点击-落地页浏览转化率"],
        "rate_lp_to_atc": ["落地页浏览-加购转化率"],
        "rate_atc_to_ic": ["加购-结账转化率"],
        "rate_ic_to_pur": ["结账-购买转化率"]
    },
    "异常指标": {
        "anomaly_metric_name": ["异常指标"],
        "mom_change": ["环比"]
    },
    "广告架构": {**COMMON_METRICS, "dimension_item": ["广告类型"]},
    "受众组": {
        **COMMON_METRICS,
        "dimension_item": ["广告组", "广告组Id", "Ad Set Name"],
        "custom_audience_settings": ["设置的自定义受众", "Custom Audiences"],
        "converting_keywords": ["产生成效的关键词", "Interests", "Keywords"],
        "converting_countries": ["产生成效的国家", "国家", "地区", "Country", "Region", "Location"],
        "converting_genders": ["产生成效的性别", "性别", "Gender"],
        "converting_ages": ["产生成效的年龄", "年龄", "Age", "Age Group"]
    },
    "受众类型": {**COMMON_METRICS, "dimension_item": ["受众类型"]},
    "国家": {**COMMON_METRICS, "dimension_item": ["国家/地区", "国家"]},
    "年龄": {**COMMON_METRICS, "dimension_item": ["年龄"]},
    "性别": {**COMMON_METRICS, "dimension_item": ["性别"]},
    "平台&版位": {**COMMON_METRICS, "dimension_item": ["平台&版位"]},
    "素材": {
        **COMMON_METRICS,
        "content_item": ["素材", "Ad Name", "Creative Name"],
        "cvr_lp_to_pur": ["落地页浏览-购买转化率"]
    },
    "落地页": {
        **COMMON_METRICS,
        "content_item": ["落地页url", "落地页", "Website URL"],
        "ctr_all": ["曝光-点击转化率"],
        "rate_lp_to_atc": ["落地页浏览-加购转化率", "落地页浏览-购物转化率"]
    }
}

GROUP_CONFIG = {
    "Master_Overview": ["整体数据", "分时段数据", "异常指标"],
    "Master_Breakdown": ["广告架构", "受众组", "受众类型", "国家", "年龄", "性别", "平台&版位"],
    "Master_Creative": ["素材", "落地页"]
}

REPORT_MAPPING = {
    "spend": "花费 ($)", "roas": "ROAS", "purchases": "购买次数", "purchase_value": "购买总价值",
    "cpa": "CPA ($)", "ctr": "CTR (%)", "cpm": "CPM ($)", "aov": "客单价",
    "impressions": "展现量", "clicks_all": "点击量 (All)", "clicks": "点击量 (All)", "ctr_all": "点击率 (All)",
    "landing_page_views": "落地页访问量", "add_to_cart": "加购次数", "initiate_checkout": "结账发起数 (IC)",
    "rate_click_to_lp": "点击 → 落地页访问转化率", "rate_lp_to_atc": "落地页 → 加购转化率",
    "rate_atc_to_ic": "加购 → 购买转化率", "rate_ic_to_pur": "购买转化率",
    "cvr_purchase": "点击 → 购买转化率", "cvr_lp_to_pur": "CVR (全站转化率)",
    "date_range": "日期/时段", "campaign_type": "投放模式", "adset_name": "广告组ID", "adset_id": "广告组ID",
    "custom_audience_settings": "自定义受众源", "converting_keywords": "高潜兴趣词", "audience_type": "受众策略",
    "country": "国家", "age_group": "年龄", "gender": "性别", "creative_name": "素材名称", "placement": "版位",
    "landing_page_url": "页面 URL", "mom_change": "环比波动", "anomaly_metric_name": "异常项",
    "converting_countries": "产生成效的国家", "converting_genders": "产生成效的性别", "converting_ages": "产生成效的年龄"
}

FIELD_ALIASES = {
    "adset_id": ["adset_id", "ad set id", "adset id", "广告组编号", "广告组id", "adset_name", "ad set name"],
    "converting_countries": ["converting_countries", "country", "region", "国家", "地区", "location"],
    "converting_genders": ["converting_genders", "gender", "性别"],
    "converting_ages": ["converting_ages", "age", "年龄", "age_group"],
    "converting_keywords": ["converting_keywords", "keywords", "interests", "兴趣", "关键词"],
    "spend": ["spend", "amount spent", "cost", "花费", "消耗"],
    "purchases": ["purchases", "results", "result", "成效", "购买"],
    "roas": ["roas", "return on ad spend", "purchase roas"],
    "purchase_value": ["purchase_value", "conversion value", "value", "总价值", "gmv", "购买总价值"],
    "clicks": ["clicks", "clicks (all)", "点击量", "clicks_all"],
    "impressions": ["impressions", "展示", "展现"],
    "ctr_all": ["ctr_all", "ctr (all)", "点击率 (all)"],
    # ✅ 修改点：增加 "网站加购", "adds to cart" 以防万一
    "add_to_cart": ["add_to_cart", "加入购物车", "加购", "cart", "website adds to cart", "网站加购", "adds to cart"], 
    "initiate_checkout": ["initiate_checkout", "结账发起次数", "结账", "checkout"],
    "landing_page_views": ["landing_page_views", "落地页浏览量", "落地页", "landing"]
}

# ==========================================
# PART 2: 核心工具函数
# ==========================================

def parse_float(value):
    if value is None: return 0.0
    try:
        if isinstance(value, (int, float)): return float(value)
        return clean_numeric_strict(value)
    except: return 0.0

def safe_div(numerator, denominator, multiplier=1.0):
    n = parse_float(numerator)
    d = parse_float(denominator)
    if d > 0: return (n / d) * multiplier
    else: return 0.0

def clean_numeric(val):
    if pd.isna(val): return 0.0
    if isinstance(val, (int, float)): return float(val)
    val_str = str(val).strip().replace('$', '').replace('¥', '').replace(',', '')
    if '%' in val_str: 
        val_str = val_str.replace('%', '')
        try: return float(val_str) / 100.0 
        except: return 0.0
    try: return float(val_str)
    except: return val # Return original if not number (for text columns)

def clean_numeric_strict(val): 
    if pd.isna(val): return 0.0
    if isinstance(val, (int, float)): return float(val)
    val_str = str(val).strip().replace('$', '').replace('¥', '').replace(',', '')
    if '%' in val_str: 
        val_str = val_str.replace('%', '')
        try: return float(val_str) / 100.0
        except: return 0.0
    try: return float(val_str)
    except: return 0.0

def find_column_fuzzy(df, keywords):
    # 1. 精确匹配
    for kw in keywords:
        if kw in df.columns: return kw
    
    # 2. 归一化匹配 (去空格、转小写)
    df_cols_norm = {c.lower().replace(' ', '').replace('_', ''): c for c in df.columns}
    for kw in keywords:
        kw_norm = kw.lower().replace(' ', '').replace('_', '')
        if kw_norm in df_cols_norm: return df_cols_norm[kw_norm]
    
    # 3. 包含匹配 (Contains)
    for col in df.columns:
        col_lower = col.lower()
        for kw in keywords:
            if kw.lower() in col_lower: return col
    return None

def calc_metrics_dict(df_chunk):
    res = {}
    if df_chunk.empty: return res
    sums = {}
    # 确保这里包含 add_to_cart
    targets = ['spend', 'clicks', 'impressions', 'purchases', 'purchase_value',
               'landing_page_views', 'add_to_cart', 'initiate_checkout']
    
    for t in targets:
        aliases = FIELD_ALIASES.get(t, [t])
        if t == 'purchase_value' and 'value' not in aliases: aliases.append('value')
        col = find_column_fuzzy(df_chunk, aliases)
        if col:
             # 直接读取列值并求和 (对于单行就是直接读取)
             sums[t] = df_chunk[col].apply(clean_numeric_strict).sum()
        else:
             sums[t] = 0.0

    res['spend'] = parse_float(sums.get('spend', 0))
    res['impressions'] = parse_float(sums.get('impressions', 0))
    res['clicks'] = parse_float(sums.get('clicks', 0))
    res['purchases'] = parse_float(sums.get('purchases', 0))
    res['purchase_value'] = parse_float(sums.get('purchase_value', 0))
    # ✅ 这里直接读取，不进行公式计算
    res['add_to_cart'] = parse_float(sums.get('add_to_cart', 0))
    res['initiate_checkout'] = parse_float(sums.get('initiate_checkout', 0))
    res['landing_page_views'] = parse_float(sums.get('landing_page_views', 0))
    
    res['roas'] = safe_div(sums.get('purchase_value'), sums.get('spend'))
    res['cpm'] = safe_div(sums.get('spend'), sums.get('impressions'), multiplier=1000)
    res['cpc'] = safe_div(sums.get('spend'), sums.get('clicks'))
    res['ctr'] = safe_div(sums.get('clicks'), sums.get('impressions'))
    res['cpa'] = safe_div(sums.get('spend'), sums.get('purchases'))
    res['cvr_purchase'] = safe_div(sums.get('purchases'), sums.get('clicks'))
    
    res['rate_click_to_lp'] = safe_div(sums.get('landing_page_views'), sums.get('clicks'))
    res['rate_lp_to_atc']   = safe_div(sums.get('add_to_cart'), sums.get('landing_page_views'))
    res['rate_atc_to_ic']   = safe_div(sums.get('initiate_checkout'), sums.get('add_to_cart'))
    res['rate_ic_to_pur']   = safe_div(sums.get('purchases'), sums.get('initiate_checkout'))
    
    res['aov'] = safe_div(sums.get('purchase_value'), sums.get('purchases'))

    date_col = find_column_fuzzy(df_chunk, ['date', 'time', 'range'])
    if date_col:
        try:
            dates = pd.to_datetime(df_chunk[date_col], errors='coerce').dropna()
            if not dates.empty: res['date_range'] = f"{dates.min():%Y-%m-%d} ~ {dates.max():%Y-%m-%d}"
            else: res['date_range'] = "-"
        except: res['date_range'] = "-"
    else: res['date_range'] = "-"
    return res 

def format_cell(key, val, is_mom=False):
    if isinstance(val, str): return val
    if is_mom:
        if key == 'date_range': return val
        return f"{val:+.2%}"
    k = str(key).lower()
    if 'roas' in k: return f"{val:.2f}"
    if any(x in k for x in ['rate', 'ctr', 'cvr', '点击率', '转化率', '着陆率', '意向率', '成功率']): 
        return f"{val:.2%}" 
    if any(x in k for x in ['spend', 'cpm', 'cpc', 'value', 'aov', 'cpa', '花费', '金额', '客单价', 'gmv', '价值']): return f"{val:,.2f}"
    if any(x in k for x in ['purchases', 'cart', 'click', '次数', '单量', '点击', '展现', '访问量', '发起数']): return f"{val:,.0f}"
    return f"{val}"

def extract_benchmark_values(df_bench):
    targets = {'roas': (['roas'], True), 'cpm': (['cpm'], False), 'ctr': (['ctr'], True), 'cpc': (['cpc'], False), 'cpa': (['cpa_purchase', 'cpa'], False)}
    extracted = {}
    for metric, (aliases, higher_better) in targets.items():
        found_col = None
        for alias in aliases:
            found_col = find_column_fuzzy(df_bench, [alias])
            if found_col: break
        if found_col:
            try:
                s = df_bench[found_col].apply(clean_numeric_strict)
                v = s[s>0].mean()
                if metric in ['ctr'] and v > 1.0: v = v / 100.0
                if not pd.isna(v): extracted[metric] = [v, higher_better]
            except: pass
    return extracted

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    try:
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if color:
            c = OxmlElement('w:color')
            c.set(qn('w:val'), color)
            rPr.append(c)
        if underline:
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink
    except: return None

def apply_report_labels(df, custom_mapping=None):
    if df.empty: return df
    mapping = REPORT_MAPPING.copy()
    if custom_mapping: mapping.update(custom_mapping)
    return df.rename(columns=mapping)

def add_df_to_word(doc, df, title, level=1):
    if df.empty: return
    doc.add_heading(title, level=level)
    t = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
    t.style = 'Table Grid'
    is_creative = "素材" in title
    is_landing = "落地页" in title
    link_col_idx = -1
    for j, col in enumerate(df.columns):
        cell = t.cell(0, j)
        cell.text = str(col)
        if any(x in str(col).lower() for x in ["url", "link", "素材", "内容", "content"]): link_col_idx = j
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(8)
    for i in range(df.shape[0]):
        label_prefix = "素材" if is_creative else ("落地页" if is_landing else "")
        label_char = chr(65 + (i % 26))
        if i >= 26: label_char += str(i // 26)
        label_text = f"{label_prefix}{label_char}"
        for j in range(df.shape[1]):
            val = df.iat[i, j]
            cell = t.cell(i+1, j)
            if (is_creative or is_landing) and j == link_col_idx:
                try:
                    p = cell.paragraphs[0]
                    url = str(val).strip()
                    if len(url) > 5: add_hyperlink(p, url, label_text)
                    else: cell.text = label_text
                except: cell.text = label_text
            else:
                cell.text = str(val)
                if "结论" in str(df.columns[j]):
                    if "✅" in str(val): cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                    if "⚠️" in str(val): cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(8)
    doc.add_paragraph("\n")

# ==========================================
# PART 3: 主逻辑类
# ==========================================

class AdReportProcessor:
    def __init__(self, raw_file, bench_file=None):
        self.raw_file = raw_file
        self.bench_file = bench_file
        self.processed_dfs = {}
        self.merged_dfs = {}
        self.final_json = {}
        self.doc = Document()

    def find_sheet_fuzzy(self, target, actual_sheets):
        for actual in actual_sheets:
            if target.strip().lower() == actual.strip().lower():
                return actual
        for actual in actual_sheets:
            if target in actual:
                return actual
        return None

    def process_etl(self):
        xls = pd.ExcelFile(self.raw_file)
        
        for config_sheet_name, mapping in SHEET_MAPPINGS.items():
            actual_sheet_name = self.find_sheet_fuzzy(config_sheet_name, xls.sheet_names)
            
            if actual_sheet_name:
                df = pd.read_excel(xls, sheet_name=actual_sheet_name)
                # 归一化列名，方便匹配
                df.columns = [str(c).strip() for c in df.columns]
                
                final_cols = {}
                for std_col, raw_col_options in mapping.items():
                    matched_col = None
                    # 1. 精确/Case-Insensitive 匹配
                    for option in raw_col_options:
                        # 查找原始列中是否存在该别名 (忽略大小写)
                        for raw_col in df.columns:
                            if option.lower() == raw_col.lower():
                                matched_col = raw_col
                                break
                        if matched_col: break
                        
                        # 如果还没找到，尝试去空格匹配
                        if not matched_col:
                            for raw_col in df.columns:
                                if option.lower().replace(" ", "") == raw_col.lower().replace(" ", ""):
                                    matched_col = raw_col
                                    break
                        if matched_col: break
                    
                    if matched_col: 
                        final_cols[std_col] = matched_col
                
                # 创建清洗后的 DataFrame
                if final_cols:
                    df_clean = df[list(final_cols.values())].rename(columns={v: k for k, v in final_cols.items()})
                else:
                    df_clean = pd.DataFrame() # 如果完全没匹配到
                
                # ✅ 核心修正：强制补全缺失的标准列，确保后续逻辑能找到 add_to_cart
                for expected_col in mapping.keys():
                    if expected_col not in df_clean.columns:
                        # 如果源文件中没找到这列，就创建它并填0
                        df_clean[expected_col] = 0.0

                # 数值清洗
                text_cols = ['date_range', 'anomaly_metric_name', 
                             'converting_keywords', 'converting_countries', 'converting_genders', 'converting_ages', 
                             'custom_audience_settings', 'dimension_item', 'content_item']
                
                for col in df_clean.columns:
                    if col not in text_cols:
                        df_clean[col] = df_clean[col].apply(clean_numeric)

                if config_sheet_name in ["素材", "落地页", "受众组"]:
                    if "spend" in df_clean.columns:
                        df_clean = df_clean.sort_values("spend", ascending=False).head(10)

                df_clean["Source_Sheet"] = config_sheet_name
                self.processed_dfs[config_sheet_name] = df_clean

        for master_name, source_sheets in GROUP_CONFIG.items():
            dfs_to_merge = [self.processed_dfs[src] for src in source_sheets if src in self.processed_dfs]
            if dfs_to_merge:
                merged_df = pd.concat(dfs_to_merge, ignore_index=True)
                cols = list(merged_df.columns)
                priority_cols = ['Source_Sheet', 'date_range', 'dimension_item', 'content_item',
                                 'spend', 'roas', 'purchases', 'cpa']
                new_order = [c for c in priority_cols if c in cols] + [c for c in cols if c not in priority_cols]
                self.merged_dfs[master_name] = merged_df[new_order]

    def generate_report(self):
        benchmark_targets = {'roas': [2.0, True], 'cpm': [20.0, False], 'ctr': [0.015, True], 'cpc': [1.5, False], 'cpa': [30.0, False]}
        if self.bench_file:
            try:
                df_b = pd.read_excel(self.bench_file)
                benchmark_targets = extract_benchmark_values(df_b)
            except: pass

        self.doc.add_heading('广告投放深度分析报告', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.final_json = {"report_title": "广告投放深度分析报告", "generated_at": pd.Timestamp.now().strftime("%Y-%m-%d")}

        # 1. 大盘总览
        df_ov = pd.DataFrame()
        if "Master_Overview" in self.merged_dfs:
            df_src = self.merged_dfs["Master_Overview"]
            mask = df_src['Source_Sheet'].astype(str).apply(lambda x: any(k in x for k in ["分时", "Time"]))
            df_ov = df_src[mask].copy() if not df_src[mask].empty else df_src.copy()

        if not df_ov.empty:
            date_col = find_column_fuzzy(df_ov, ['date', 'time', '时间'])
            if date_col:
                try:
                    df_ov['temp_date'] = pd.to_datetime(df_ov[date_col], errors='coerce')
                    df_clean = df_ov.dropna(subset=['temp_date']).sort_values('temp_date')
                    dates = df_clean['temp_date'].unique()
                    
                    # 1.1 基于分时数据的基础计算
                    raw_overall = calc_metrics_dict(df_clean)
                    
                    # ======================================================
                    # ✅ [核心逻辑修正] 覆盖数据逻辑增强
                    # ======================================================
                    if "Master_Overview" in self.merged_dfs:
                         df_all = self.merged_dfs["Master_Overview"]
                         mask_summary = df_all['Source_Sheet'] == "整体数据"
                         df_summary = df_all[mask_summary]
                         
                         if not df_summary.empty:
                             summary_row = df_summary.iloc[0]
                             override_metrics = ['add_to_cart', 'initiate_checkout', 'purchases', 'landing_page_views', 'impressions', 'clicks']
                             
                             for m in override_metrics:
                                 # 只要列存在，就尝试读取
                                 if m in summary_row:
                                     val = clean_numeric_strict(summary_row[m])
                                     # 只有值大于0才覆盖，防止坏数据
                                     if val > 0:
                                         raw_overall[m] = val
                             
                             # 🚨 重新计算转化率 (因为分子分母变了)
                             raw_overall['rate_click_to_lp'] = safe_div(raw_overall.get('landing_page_views'), raw_overall.get('clicks'))
                             raw_overall['rate_lp_to_atc']   = safe_div(raw_overall.get('add_to_cart'), raw_overall.get('landing_page_views'))
                             raw_overall['rate_atc_to_ic']   = safe_div(raw_overall.get('initiate_checkout'), raw_overall.get('add_to_cart'))
                             raw_overall['rate_ic_to_pur']   = safe_div(raw_overall.get('purchases'), raw_overall.get('initiate_checkout'))
                             raw_overall['cvr_purchase'] = safe_div(raw_overall.get('purchases'), raw_overall.get('clicks'))
                    # ======================================================

                    if len(dates) >= 2:
                        mid_date = dates[len(dates)//2]
                        raw_prev = calc_metrics_dict(df_clean[df_clean['temp_date'] < mid_date])
                        raw_curr = calc_metrics_dict(df_clean[df_clean['temp_date'] >= mid_date])
                        raw_mom = {}
                        for k, v_curr in raw_curr.items():
                            if k == 'date_range': raw_mom[k] = "-"
                            else:
                                v_prev = raw_prev.get(k, 0)
                                raw_mom[k] = (v_curr - v_prev) / v_prev if v_prev > 0 else 0.0
                    else:
                        raw_prev = {k: "-" for k in raw_overall}; raw_curr = raw_overall; raw_mom = {k: "-" for k in raw_overall}

                    col_order = ["date_range", "spend", "roas", "cpa", "cpm", "cpc", "ctr", "cvr_purchase",
                                 "rate_click_to_lp", "rate_lp_to_atc", "rate_ic_to_pur", "aov", "add_to_cart", "purchases", "purchase_value"]
                    final_data = []
                    for label, r in zip(["整体数据", "上周期值", "本周期", "环比"], [raw_overall, raw_prev, raw_curr, raw_mom]):
                        row = {"Label": label}
                        is_m = (label == "环比")
                        for c in col_order: row[c] = format_cell(c, r.get(c, 0), is_mom=is_m)
                        row['date_range'] = label
                        final_data.append(row)

                    df_f = pd.DataFrame(final_data, columns=col_order)
                    df_f_display = apply_report_labels(df_f)
                    add_df_to_word(self.doc, df_f_display, "1. 数据大盘总览", level=1)
                    self.final_json['1_data_overview'] = df_f_display.to_dict(orient='records')

                    # 2. Benchmark
                    raw_current = raw_overall.copy()
                    bench_data = []
                    for metric_key in ['roas', 'cpm', 'ctr', 'cpc', 'cpa']:
                        curr_val = raw_current.get(metric_key, 0)
                        bench_val, higher_is_better = benchmark_targets.get(metric_key, [0, True])
                        conclusion = "-"
                        if curr_val != 0:
                            diff = curr_val - bench_val
                            if higher_is_better: conclusion = "✅ 优于大盘" if diff > 0 else ("⚠️ 低于大盘" if diff < 0 else "持平")
                            else: conclusion = "✅ 优于大盘" if diff < 0 else ("⚠️ 高于大盘" if diff > 0 else "持平")
                        bench_data.append({
                            "指标": REPORT_MAPPING.get(metric_key, metric_key.upper()),
                            "当前账户": format_cell(metric_key, curr_val),
                            "行业基准": format_cell(metric_key, bench_val),
                            "对比结论": conclusion
                        })
                    df_b = pd.DataFrame(bench_data)
                    add_df_to_word(self.doc, df_b, "2. 行业 Benchmark 对比", level=1)
                    self.final_json['2_industry_benchmark'] = df_b.to_dict(orient='records')
                except Exception as e: st.warning(f"大盘计算警告: {e}")

        # 3. 受众组
        self.generate_audience_section()
        # 4. 素材与落地页
        self.generate_creative_section()
        # 5. 版位
        self.generate_placement_section()
        # 7. 架构诊断
        self.generate_structure_section()

    def generate_audience_section(self):
        self.doc.add_heading("3. 受众组分析", level=1)
        self.final_json['3_audience_analysis'] = {}
        audience_configs = [
            ("3.1 国家分析", ["国家", "Country"], True, "国家"),
            ("3.2 性别分析", ["性别", "Gender"], False, "性别"),
            ("3.3 年龄分析", ["年龄", "Age"], False, "年龄段"),
            ("3.4 受众组分析表", ["受众", "Audience"], True, "受众组名称"),
        ]
        if "Master_Breakdown" in self.merged_dfs:
            df_bd = self.merged_dfs["Master_Breakdown"]
            for title, keywords, top10, dim_label in audience_configs:
                mask = df_bd['Source_Sheet'].astype(str).apply(lambda x: any(k in x for k in keywords))
                df_curr = df_bd[mask].copy()
                if not df_curr.empty:
                    self.process_sub_table(df_curr, title, top10, dim_label, '3_audience_analysis')

    def generate_creative_section(self):
        if "Master_Creative" in self.merged_dfs:
            df_cr = self.merged_dfs["Master_Creative"]
            for title, keywords, label, json_key in [("4. 素材分析", ["素材", "Creative"], "素材名称", "4_creative_analysis"), ("6. 落地页分析", ["落地页", "Landing"], "落地页 URL", "6_landing_page_analysis")]:
                mask = df_cr['Source_Sheet'].astype(str).apply(lambda x: any(k in x for k in keywords))
                df_curr = df_cr[mask].copy()
                if not df_curr.empty:
                      # 简单的CPC/CTR补全逻辑，同原代码
                      if not find_column_fuzzy(df_curr, ['cpc']): df_curr['cpc'] = df_curr['spend'] / df_curr['clicks'].replace(0, np.nan) if 'clicks' in df_curr else 0
                      if not find_column_fuzzy(df_curr, ['cpa']): df_curr['cpa'] = df_curr['spend'] / df_curr['purchases'].replace(0, np.nan) if 'purchases' in df_curr else 0
                      if not find_column_fuzzy(df_curr, ['ctr']): df_curr['ctr'] = (df_curr['clicks'] / df_curr['impressions'].replace(0, np.nan)) * 100 if 'impressions' in df_curr else 0
                      else: df_curr['ctr'] = df_curr['ctr'] * 100
                      
                      req_cols = ["content_item", "spend", "ctr", "cpc", "cpm", "roas", "cpa"]
                      df_final = self.standardize_cols(df_curr, req_cols)
                      if 'spend' in df_final.columns: df_final = df_final.sort_values('spend', ascending=False).head(10)
                      df_display = apply_report_labels(df_final.round(2), custom_mapping={'content_item': label})
                      add_df_to_word(self.doc, df_display, title, level=1)
                      self.final_json[json_key] = df_display.to_dict(orient='records')

    def generate_placement_section(self):
         if "Master_Breakdown" in self.merged_dfs:
             self.doc.add_heading("5. 版位分析", level=1)
             df_bd = self.merged_dfs["Master_Breakdown"]
             mask = df_bd['Source_Sheet'].astype(str).apply(lambda x: any(k in x for k in ["版位", "Placement"]))
             df_curr = df_bd[mask].copy()
             if not df_curr.empty:
                  req_cols = ['dimension_item', 'spend', 'ctr', 'cpc', 'cpm', 'roas', 'cpa']
                  # 简单补全计算
                  if 'clicks' in df_curr and 'impressions' in df_curr: df_curr['ctr'] = df_curr['clicks'] / df_curr['impressions'].replace(0,np.nan)
                  
                  df_clean = self.standardize_cols(df_curr, req_cols).round(2)
                  df_top5 = df_clean.sort_values('spend', ascending=False).head(5)
                  add_df_to_word(self.doc, apply_report_labels(df_top5, {'dimension_item': '版位'}), "5.1 版位花费 TOP 5", level=2)
                  self.final_json['5_placement_analysis'] = {"top_spend": df_top5.to_dict('records')}

    def generate_structure_section(self):
        rows = []
        if "Master_Overview" in self.merged_dfs:
             metrics = calc_metrics_dict(self.merged_dfs["Master_Overview"])
             rows.append({"模块": "预算结构", "当前结构数据表现": f"总花费: ${metrics.get('spend',0):,.2f}", "存在的问题": ""})
        df_struct = pd.DataFrame(rows)
        add_df_to_word(self.doc, df_struct, "7. 广告架构分析", level=1)
        self.final_json['7_structure_analysis'] = df_struct.to_dict(orient='records')

    def standardize_cols(self, df, req_cols):
        rename_map = {}; valid_cols = []
        for req in req_cols:
            aliases = FIELD_ALIASES.get(req, [req])
            found = find_column_fuzzy(df, aliases)
            if found: valid_cols.append(found); rename_map[found] = req
            else: df[req] = 0.0; valid_cols.append(req)
        return df[valid_cols].rename(columns=rename_map)

    def process_sub_table(self, df, title, top10, dim_label, json_section):
        req_cols = ["dimension_item", "spend", "ctr", "cpc", "cpm", "cpa", "roas"]
        if "受众" in title: req_cols += ["converting_countries", "converting_keywords"]
        df_final = self.standardize_cols(df, req_cols)
        if top10 and 'spend' in df_final.columns: df_final = df_final.sort_values('spend', ascending=False).head(10)
        df_display = apply_report_labels(df_final.round(2), custom_mapping={'dimension_item': dim_label})
        add_df_to_word(self.doc, df_display, title, level=2)
        if json_section not in self.final_json: self.final_json[json_section] = {}
        self.final_json[json_section][title] = df_display.to_dict(orient='records')

# ==========================================
# PART 4: Streamlit UI
# ==========================================
def main():
    st.set_page_config(page_title="Auto-ad-data", layout="wide")
    st.title("广告数据自动化清洗系统")
    
    raw_file = st.file_uploader("1.上传【周期性复盘报告】", type=["xlsx", "xls"])
    bench_file = st.file_uploader("2.上传【行业 Benchmark】", type=["xlsx", "xls"])
    
    if st.button("开始生成数据表") and raw_file:
        processor = AdReportProcessor(raw_file, bench_file)
        try:
            with st.spinner("数据处理中..."):
                processor.process_etl()
                processor.generate_report()
            st.success("处理完成！")
            
            # 下载按钮逻辑
            json_str = json.dumps(processor.final_json, indent=4, ensure_ascii=False)
            st.download_button("📥 下载 JSON", json_str, "report.json", "application/json")
            
            output_doc = io.BytesIO()
            processor.doc.save(output_doc)
            st.download_button("📥 下载 Word", output_doc.getvalue(), "report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
        except Exception as e:
            st.error(f"发生错误: {str(e)}")
            st.exception(e)

if __name__ == "__main__":
    main()
